import requests
import openpyxl
import csv
import tkinter as tk
import os
from bs4 import BeautifulSoup
from urllib.parse import urljoin, unquote
from tkinter import messagebox, filedialog
from tkinter.ttk import Progressbar
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Inches, Pt
from docx import Document as WordDocument
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from threading import Thread



app = None 

# Obtain the final url from the relative url
def get_final_url(url):
    try:
        response = requests.get(url, allow_redirects=True)
        return response.url
    except requests.exceptions.RequestException:
        return None

# Clean the link, don't check mailto or anchors
def clean_link(link):
    if link and not link.startswith("mailto:") and not link.startswith("#"):
        link = unquote(link).strip()
        # CHANGED TO WHILE LOOP
        while link.endswith("%20"):
            link = link[:-1]
        return link
    return None

# Get all the links from page
def get_links_from_page(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            links = [clean_link(link.get('href')) for link in soup.find_all('a', href=True)]
            return [link for link in links if link is not None]  # Filter out None values
        else:
            return []
    except requests.exceptions.RequestException:
        return []

# Identify whether the link is on navigation, footer, or main content (Add left navigation)
def identify_section(link, page_link):
    try:
        absolute_page_link = urljoin(link, page_link)
        response = requests.get(link)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            # Search for partial matches of the href attribute with page_link
            element = soup.find('a', href=lambda href: page_link in href)
            if element:
                # Search for the outer elements to determine the section
                # change to id siteheader
                nav_element = element.find_parent('nav')
                if nav_element:
                    return "Navigation"

                footer_element = element.find_parent('footer')
                if footer_element:
                    return "Footer"

                sidr_container_element = element.find_parent('div', id='sidr-container')
                if sidr_container_element:
                    return "Main Content"

            # BELOW SHOULDN'T BE USED, JUST IN CASE
            return "Unknown"  # If the element is not found in any of the relevant outer elements
        else:
            return "Unknown"  # If there was an issue fetching the page
    except requests.exceptions.RequestException:
        return "Unknown"  # If there was a connection error or other issues

# Main function to check for broken links given an excel
def check_broken_links(excel_file):
    global app
    try:
        print(f"Opening file: {excel_file}...")
        wb = openpyxl.load_workbook(excel_file)
        sheet = wb.active
        column_with_links = 'A'

        broken_links_report = []
        total_links_checked = 0
        total_links = 0

        for cell in sheet[column_with_links]:
            if cell.value:
                total_links += 1

        print(f"Checking {total_links} links inside {excel_file}...")

        # Create a Progressbar widget for the overall progress
        progress_bar = Progressbar(app.root, orient='horizontal', length=300, mode='determinate')
        progress_bar.pack()

        for cell in sheet[column_with_links]:
            if cell.value:
                link = cell.value
                total_links_checked += 1
                page_links = get_links_from_page(link)

                # Page progress bar starts here
                page_progress_var = tk.DoubleVar()
                page_progress = Progressbar(app.root, variable=page_progress_var, length=300, maximum=len(page_links),
                                            mode='determinate')
                page_progress.pack()

                # Update page progress label
                app.update_page_progress(total_links_checked, total_links)

                for i, page_link in enumerate(page_links):
                    absolute_link = urljoin(link, page_link)
                    final_url = get_final_url(absolute_link)
                    section = identify_section(link, page_link)

                    if final_url is not None:
                        try:
                            response = requests.get(final_url)
                            if response.status_code != 200:
                                if (link, final_url, section, page_link) not in broken_links_report:  # Check for duplicates
                                    broken_links_report.append((link, final_url, section, page_link))
                                    print(f"Broken link found on page '{link}' in '{section}' section: {final_url}",
                                          end='\r')
                                else:
                                    print(f"Duplicate broken link found on page '{link}': {final_url}", end='\r')
                            else:
                                pass  # No need for this message
                        except requests.exceptions.RequestException:
                            if (link, final_url, section, page_link) not in broken_links_report:  # Check for duplicates
                                broken_links_report.append((link, final_url, section, page_link))
                                print(
                                    f"Connection error occurred for link on page '{link}' in '{section}' section: {final_url}",
                                    end='\r')
                            else:
                                print(f"Duplicate connection error occurred for link on page '{link}': {final_url}",
                                      end='\r')
                    else:
                        # Include invalid links in the broken links report
                        broken_links_report.append((link, page_link, section, "Invalid URL"))
                        print(f"Invalid URL: {absolute_link}", end='\r')
                    page_progress_var.set(i + 1)
                    app.update_link_progress(i + 1, len(page_links))
                    app.root.update_idletasks()

                page_progress.destroy()

            progress_bar['value'] = total_links_checked
            app.root.update_idletasks()

        print("\nLink checking completed.")
        print(f"Total pages checked: {total_links_checked}")
        print(f"Total broken links found: {len(broken_links_report)}")
        progress_bar.destroy()
        return broken_links_report
    except openpyxl.utils.exceptions.InvalidFileException:
        print(f"Error: Unable to open the file '{excel_file}'. Please check if it's a valid Excel file.")
        return []
    except Exception as e:
        print(f"Error: An unexpected error occurred while processing the file '{excel_file}': {e}")
        return []

def save_to_csv(file_path, data, header):
    with open(file_path, 'w', newline='') as csvfile:
        csv_writer = csv.writer(csvfile)
        csv_writer.writerow(header)
        csv_writer.writerows(data)

def create_word_document_from_excel(broken_links_report, output_file):
    document = WordDocument()

    # Add a title to the document
    document.add_heading("Broken Links Report", level=1)

    # Iterate through the broken links and add content to the document
    for link_info in broken_links_report:
        link, page_link, section, relative = link_info
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"• On {link} there is a broken link on section '{section}', the broken link is '{page_link}'.")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

        # Add a bullet point after each link
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(0.5)  # Adjust the left indentation for the bullet point
        paragraph_format.space_after = Pt(12)  # Add spacing after each paragraph

    # Save the document to the specified file path
    document.save(output_file)

def get_folders_in_same_level():
    script_directory = os.path.dirname(os.path.abspath(__file__))
    return [folder for folder in os.listdir(script_directory) if os.path.isdir(os.path.join(script_directory, folder))]


# Get the list of Excel files in the selected folder
def get_excel_files(folder_path):
    return [f for f in os.listdir(folder_path) if f.endswith(".xlsx")]

def update_excel_list(app_instance):
    global app  # Use the global app variable
    app = app_instance  # Set the global app variable to the app instance

    selected_folder = app.folder_var.get()
    excel_files = get_excel_files(selected_folder)
    app.excel_listbox.delete(*app.excel_listbox.get_children())  # Clear the existing list

    # Insert the Excel files into the Treeview widget
    for idx, excel_file in enumerate(excel_files, start=1):
        app.excel_listbox.insert("", "end", values=(excel_file,))

    print(f"Found {len(excel_files)} Excel file(s) in folder: {selected_folder}")


def save_report_to_folder(folder_path, excel_name, report_data):
    # Create a folder with the name of the Excel file (if it doesn't exist)
    excel_folder = os.path.join(folder_path, excel_name)
    if not os.path.exists(excel_folder):
        os.makedirs(excel_folder)

    # Check if the excel_folder is a valid directory
    if not os.path.isdir(excel_folder):
        raise NotADirectoryError(f"'{excel_folder}' is not a valid directory.")

    # Get the number of existing reports in the folder
    report_count = len([f for f in os.listdir(excel_folder) if f.startswith("Report_")])

    # Create a folder for the current report with the format "Report_x_{time}"
    now = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_folder_name = f"Report_{report_count + 1}_{now}"
    report_folder = os.path.join(excel_folder, report_folder_name)
    os.makedirs(report_folder)

    # Save the broken links report to a CSV file inside the report folder
    csv_file_path = os.path.join(report_folder, "broken_links_report.csv")
    header = ["Page", "Broken Link", "Section"]
    save_to_csv(csv_file_path, report_data, header)

    return report_folder