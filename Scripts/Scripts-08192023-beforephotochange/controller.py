import requests
import csv
import os
from bs4 import BeautifulSoup
from urllib.parse import urljoin, unquote
from docx.shared import RGBColor, Inches, Pt
from docx import Document as WordDocument
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from concurrent.futures import ThreadPoolExecutor
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException
from PIL import Image
import io
import base64
import time
import urllib.parse
import urllib

class Controller:
    def __init__(self):
        self.controllerCreated = True
        
    def get_final_url(self,url):
        try:
            response = requests.get(url, allow_redirects=True, timeout=10)
            return response.url
        except requests.exceptions.RequestException:
            return None

        
    def clean_link(self, link):
        if link and not link.startswith("mailto:") and not link.startswith("#") and not link.startswith("tel:"):
            link = unquote(link).strip()
            # CHANGED TO WHILE LOOP
            while link.endswith("%20"):
                link = link[:-1]
            return link
        return None


    def get_links_from_page_concurrently(self, url):
        try:
            with ThreadPoolExecutor() as executor:
                response = requests.get(url)
                if response.status_code == 200:
                    soup = BeautifulSoup(response.content, 'html.parser')
                    links = [self.clean_link(link.get('href')) for link in soup.find_all('a', href=True)]
                    return [link for link in links if link is not None]  # Filter out None values
                else:
                    return []
        except requests.exceptions.RequestException:
            return []
    
    def check_link_status(self, link):
        try:
            response = requests.get(link)
            if response.status_code != 200:
                return link, response.url
            return None
        except requests.exceptions.RequestException:
            return link, "Connection error"
    
    def save_to_csv(self, file_path, data, header):
        with open(file_path, 'w', newline='') as csvfile:
            csv_writer = csv.writer(csvfile)
            csv_writer.writerow(header)
            csv_writer.writerows(data)

    def create_word_document_from_excel(self, broken_links_report, output_file):
        document = WordDocument()

        # Add a title to the document
        document.add_heading("Broken Links Report", level=1)

        # Iterate through the broken links and add content to the document
        for link_info in broken_links_report:

            link, page_link, section, relative, link_name = link_info
            
            try:
                photo = self.get_photo(link, link_name, relative)
            except Exception as e:
                photo='icon/photo_no_available.png'
            
            if 'None' in photo:
                photo='icon/photo_no_available.png'
            paragraph = document.add_paragraph()
            link_name=self.word_link_cleaner(link_name)
            # Add the link as a run with hyperlink properties
            run = paragraph.add_run("• Broken link on ")
            self.add_hyperlink(paragraph,link,link)
                
            # Add the rest of the text
            run = paragraph.add_run(f" on ")
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
            
            # Add bold section text
            run = paragraph.add_run(section)
            run.bold = True
            
            run = paragraph.add_run(f" section, the broken link is displayed as: ")
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
            
            # Add bold section text
            run = paragraph.add_run(link_name)
            run.bold = True
            
            run = paragraph.add_run(f" and the link is pointing to : ")
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

            # Add bold link_name text
            self.add_hyperlink(paragraph, page_link, page_link)


            #PUT PHOTO HERE ON WORD
            if photo:  # Check if photo is obtained
                document.add_picture(photo, width=Inches(5))  # Add photo to document


            # Add a bullet point after each link
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0.5)  # Adjust the left indentation for the bullet point
            paragraph_format.space_after = Pt(12)  # Add spacing after each paragraph

        # Provide the path of the folder you want to clean
        self.clean_folder('photos')

        # Save the document to the specified file path
        document.save(output_file)  

    def word_link_cleaner(self, link):
        if 'http' in link:
            return link
        else:
            try:
                links = link.split('/')
            except Exception as e:
                pass
            return links[-1]      

    def get_photo(self, link, link_name, relative_link):
        # Initialize WebDriver and navigate to the webpage
        driver = webdriver.Chrome()
        driver.maximize_window()  # Open the WebDriver in full screen mode

        driver.get(f"{link}")



        # Search for the sequence of tabs/dropdowns
        search_sequence = link_name
        print(f'Getting photo, link name is: {link_name}')
        # Check the conditions before splitting
        if not search_sequence.startswith('http') and search_sequence.count('/') < 3:
            clicks = search_sequence.split('/')
        else:
            clicks = [search_sequence]  # Keep the entire sequence as a single step
        search_word = clicks[-1]  # The last element in the sequence

        #target_href = "accordion3"
        if "'" in search_word:
            search_word = search_word.split("'")[0]

        # Clean up the search_word
        search_word = search_word.lstrip(" =-,.!?").strip()


        # Loop through the clicks, excluding the last one
        while len(clicks) > 1:  # Exclude the last element
            print('Tab or accordion detected...')
            click = clicks.pop(0)  # Get the next click from the list
            while True:
                try:
                    print(click)
                    # Find and click on the tab or dropdown element
                    if 'accordion' in click or 'tab' in click:
                        click_element = driver.find_element(By.CSS_SELECTOR, f"a[href='#{click}']")
                    else:
                        click_element = driver.find_element(By.XPATH, f"//a[contains(text(), '{click}')]")

                    
                    # Scroll the element into view using JavaScript
                    driver.execute_script("arguments[0].scrollIntoView();", click_element)

                    # Give some time for the page to adjust after scrolling
                    time.sleep(1)

                    # Scroll a bit more to ensure proper visibility
                    driver.execute_script("window.scrollBy(0, -400);")
                    click_element.click()
                    
                    break
                except NoSuchElementException:
                    print(f"Element '{click}' not found. Moving to the next step.")
                    # Get the current scroll height
                    scroll_height = driver.execute_script("return document.body.scrollHeight")

                    # Scroll down
                    driver.find_element(By.TAG_NAME, "body").send_keys(Keys.END)

                    # Give some time for content to load (you may need to adjust this)
                    time.sleep(1)

                    # Check if you've reached the end
                    if scroll_height == prev_scroll_height:
                        print("Reached the end of the page.")
                        break

                    # Update the previous scroll height
                    prev_scroll_height = scroll_height


        # Define a variable to keep track of previous scroll height
        prev_scroll_height = 0
        photo_name = search_word
        photo='photo_no_available.png'
        # Scroll loop
        while True:
            print('Taking picture')
            try:

                print(f'Search word is {search_word}')
                element = driver.find_element(By.XPATH, f"//*[contains(text(), '{search_word}')]")
                print(f'element all is: {element}')

                element = driver.find_element(By.XPATH, f"//a[contains(text(), '{search_word}')]")
                print(f'element a is: {element}')
                if 'Online' in photo_name:

                    elements1 = driver.find_elements(By.XPATH, f"//*[contains(text(), '{photo_name}')]")

                    try:
                        print(f'Relative link is: {relative_link}')
                        relative_href = relative_link.split('/')[-1]  # Extract the fragment from the full URL
                        print(f'Relative href is: {relative_href}')
                        elements2 = driver.find_elements(By.CSS_SELECTOR, f"a[href*='{relative_href}']")
                        print(f'Elements 1 are: {elements1}')
                        print(f'Elements 2 are: {elements2}')
                    except Exception as e:
                        print(f'Error is {e}')
                        elements2=[]

                    print(f'Elements 1 are: {elements1}')
                    print(f'Elements 2 are: {elements2}')

                    common_elements = set(elements1) & set(elements2)

                    print("Elements present in both lists:")
                    for element in common_elements:
                        print(element.get_attribute("outerHTML"))
                    
                #if len(click_elements) == 1:
                #    click_element = click_elements[0]
                #else:
                #    click_element = driver.find_element(By.XPATH, f"//a[contains(text(), '{click}') and @href='#{relative_link}']")
                #    print(f"The click element is then: {click_element}")
                if element is None:
                    parts = search_word.split('/')
                    common_elements = []

                    for part in parts:
                        common_elements_part = driver.find_elements(By.XPATH, f"//a[contains(@href, '{part}')]")
                        if common_elements_part:
                            if not common_elements:
                                common_elements = common_elements_part
                            else:
                                common_elements = [element for element in common_elements if element in common_elements_part]

                    if common_elements:
                        element = common_elements[0]  # Returning the first common element found
                    else:
                        element = None
                        print('Element not found in page')

                # Get location and size of element
                location = element.location
                size = element.size
                

                # Scroll a bit more to ensure proper visibility
                driver.execute_script(f"window.scrollBy(0, {location['y']-400});")
                
                # Give some time for the page to adjust after scrolling
                time.sleep(1)
                
                element = driver.find_element(By.XPATH, f"//*[contains(text(), '{search_word}')]")
                
                
                # Give some time for the page to adjust after scrolling
                time.sleep(1)
                
                # Get location and size of element
                location = element.location
                size = element.size
                
                # Define screenshot area around element
                if location['y']<200:
                    y=100
                    h=600
                else:
                    y = 150
                    h = 600
                
                screenshot_base64 = driver.get_screenshot_as_base64()
                
                # Convert the base64-encoded screenshot to a PIL Image
                screenshot_bytes = base64.b64decode(screenshot_base64)
                pil_image = Image.open(io.BytesIO(screenshot_bytes))
                
                
                pil_image.save(f'photos/{photo_name}nocrop.png')

                cropped_image = pil_image.crop((0, y, pil_image.width, y + h))  # Crop only the specified height

                
                #print(cropped_image)
                
                cropped_image.save(f"photos/{photo_name}.png")
                
                # Save the screenshot
                #pil_image.save("element_screenshot_pil1.png")
                
                print("Screenshot of the entire page captured.")
                break
            except NoSuchElementException:
                print("Element not found. Scrolling...")
            
            except Exception as e:
                print(f"The error is {e}")
                
            # Get the current scroll height
            scroll_height = driver.execute_script("return document.body.scrollHeight")
            
            # Scroll down
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.END)
            
            # Give some time for content to load (you may need to adjust this)
            time.sleep(2)
            
            # Check if you've reached the end
            if scroll_height == prev_scroll_height:
                print("Reached the end of the page.")
                break
            
            # Update the previous scroll height
            prev_scroll_height = scroll_height
            
        # Close the WebDriver
        driver.quit()

        return f"photos/{photo_name}.png"


    def get_files_in_excel_folder(self):
        excel_folder = "Documents/Excel_Files"
        return [file for file in os.listdir(excel_folder) if file.endswith(".xlsx")]

        #Identify whether the link is on navigation, footer, or main content (Add left navigation)
    def identify_section(self, link, page_link, page_style):
        try:
            absolute_page_link = urljoin(link, page_link)
            response = requests.get(link)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                # Search for partial matches of the href attribute with page_link
                page_link=urllib.parse.unquote(page_link)  # Remove any URL encoding

                delimiter_characters = ['~', '%0D', '%']
                common_elements = []

                for delimiter in delimiter_characters:
                    if delimiter in page_link:
                        print('~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~')
                        print(page_link)
                        print(delimiter)

                        parts = page_link.split(delimiter)
                        
                        print(parts)
                        elements_part = [soup.find_all('a', href=lambda href: part in href) for part in parts]
                        print(f'Element parts are: {elements_part}')
                        if elements_part:
                            if not common_elements:
                                common_elements = elements_part
                            else:
                                common_elements = [element for element in common_elements if element in elements_part]
                        print(common_elements)
                    if common_elements:
                        element = common_elements[0][0]  # Returning the first common element found
                        break
                    else:
                        element = soup.find('a', href=lambda href: page_link in href)
                
                if element is None:
                    for delimiter in delimiter_characters:
                        if delimiter in saved_link:

                            parts = saved_link.split(delimiter)

                            print(parts)
                            elements_part = [soup.find_all('a', href=lambda href: part in href) for part in parts]
                            print(f'Element parts are: {elements_part}')
                            if elements_part:
                                if not common_elements:
                                    common_elements = elements_part
                                else:
                                    common_elements = [element for element in common_elements if element in elements_part]
                        
                        if common_elements:
                            if common_elements[0]:
                                element = common_elements[0][0]  # Get the link from the second inner list
                            elif common_elements[1]:
                                element = common_elements[1][0]  # Get the link from the first inner list
                        else:
                            element = soup.find('a', href=lambda href: saved_link in href)

                if element:
                    element_name = 'No_name'
                    element_name = element.get_text(strip=True)
                    
                    # Check for accordion or tab
                    accordion_div = element.find_parent('div', class_='collapsable collapse')
                    if accordion_div:
                        accordion_id = accordion_div.get('id')
                        if accordion_id:
                            element_name = f"{accordion_id}/{element_name}"

                    tab_div = element.find_parent('div', class_='tab-pane')
                    if tab_div:
                        tab_id = tab_div.get('id')
                        if tab_id:
                            element_name = f"{tab_id}/{element_name}"


                    if page_style == "No Carousel":

                        left_nav_element = element.find_parent('div', class_='leftSidebar')
                        if left_nav_element:
                            return "Left Navigation", element_name

                        main_element = element.find_parent('div', class_='rightSidebar')
                        if main_element:
                            return "Main Content", element_name   
                        
                        breadcrumb = element.find_parent(class_='breadcrumb')
                        if breadcrumb:
                            return "Main Content (Top small navigation)", element_name
                        
                        nav_element = element.find_parent('nav')
                        if nav_element:
                            # Find the outer ul element with class "nav navbar-nav" within the nav_element
                            outer_ul = nav_element.find('ul', class_='nav navbar-nav')

                            if outer_ul and 'navbar-right' not in outer_ul.get('class', []):
                                # Loop through <li> elements with class 'dropdown'
                                for li in outer_ul.find_all('li', class_='dropdown'):
                                    # Check if the element is a descendant of the current <li>
                                    if element in li.descendants:
                                        # Find the specific <a> element within the <li>
                                        dropdown_a = li.find('a', class_='dropdown-toggle')
                                        if dropdown_a:
                                            dropdown_name = dropdown_a.get_text(strip=True)
                                            return f"Main Navigation, tab name: {dropdown_name}", f"{dropdown_name}/{element_name}"
                                    
                                #return "Main Navigation, tab name: {}", element_name

                        resources_li = element.find_parent('li', class_='dropdown resourcesForLinks')
                        if resources_li:
                            return "Main Navigation, tab name: Resources For Links", f"Resources For/{element_name}"

                        # If not found in dropdowns, check in "Quick Links"
                        quick_links_li = element.find_parent('li', class_='hidden-below-1365')
                        if quick_links_li:
                            return "Main Navigation, tab name: Quick Links", f"Quick Links/{element_name}"

                        footer_element = element.find_parent('footer')
                        if footer_element:
                            return "Footer", element_name

                        #Fullscreen
                        section_anchor_div = soup.find('div', class_='sectionAnchor')
                        if section_anchor_div:
                            section_id = section_anchor_div.get('id')
                            
                            if section_id:
                                return f"Main Content: Section ID: {section_id}" , element_name

                            else:
                                return f"Main Content: Section ID: No ID", element_name

                        
                        
                        sidr_container_element = element.find_parent('div', id='sidr-container')
                        if sidr_container_element:
                            return "Somewhere unknown", element_name
                    else:
                        return "TemplateUnknown","None"
                #Error case
                return "Unknown-notfound","None"  # If the element is not found in any of the relevant outer elements
            else:
                return "Unknown-pageerror","None"  # If there was an issue fetching the page
        except requests.exceptions.RequestException:
            return "Unknown-connectionerror","None"  # If there was a connection error or other issues

    import os

    def clean_folder(self, folder_path):
        # Check if the folder exists
        if os.path.exists(folder_path) and os.path.isdir(folder_path):
            # Iterate over all files and subdirectories in the folder
            for item in os.listdir(folder_path):
                item_path = os.path.join(folder_path, item)
                if os.path.isfile(item_path):
                    # Remove files
                    os.remove(item_path)
                elif os.path.isdir(item_path):
                    # Remove subdirectories and their contents recursively
                    clean_folder(item_path)
                    os.rmdir(item_path)
        else:
            print(f"The folder '{folder_path}' does not exist.")

    def add_hyperlink(self, paragraph, url, text):
        """
        A function that places a hyperlink within a paragraph object.
        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """
        
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        # Create a new w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.ns.qn('w:val'), 'single')  # 'single' for single underline
        rPr.append(u)

        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(docx.oxml.ns.qn('w:val'), '0000FF')  # Blue color as HEX value
        rPr.append(color)

        new_run.append(rPr)

        # Join all the xml elements together and add the required text to the w:r element
        t = docx.oxml.shared.OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)

        # Append the hyperlink to the paragraph
        paragraph._p.append(hyperlink)

        return hyperlink
