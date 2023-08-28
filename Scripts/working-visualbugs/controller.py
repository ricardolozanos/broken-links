import requests
import csv
import os
from bs4 import BeautifulSoup
from urllib.parse import urljoin, unquote
from docx.shared import RGBColor, Inches, Pt
from docx import Document as WordDocument
from concurrent.futures import ThreadPoolExecutor

class Controller:
    def __init__(self):
        self.controllerCreated = True
        
    def get_final_url(self,url):
        try:
            response = requests.get(url, allow_redirects=True)
            return response.url
        except requests.exceptions.RequestException:
            return None

        
    def clean_link(self, link):
        if link and not link.startswith("mailto:") and not link.startswith("#"):
            link = unquote(link).strip()
            # CHANGED TO WHILE LOOP
            while link.endswith("%20"):
                link = link[:-1]
            return link
        return None

    def get_links_from_page(self, url):
        try:
            response = requests.get(url)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                links = [self.clean_link(link.get('href')) for link in soup.find_all('a', href=True)]
                return [link for link in links if link is not None]  # Filter out None values
            else:
                return []
        except requests.exceptions.RequestException:
            return []

    def get_links_from_page_concurrently(self, url):
        try:
            with ThreadPoolExecutor() as executor:
                response = requests.get(url)
                if response.status_code == 200:
                    soup = BeautifulSoup(response.content, 'html.parser')
                    links = [self.clean_link(link.get('href')) for link in soup.find_all('a', href=True)]
                    return [link for link in links if link is not None]  # Filter out None values
                else:
                    return []
        except requests.exceptions.RequestException:
            return []
    
    def check_link_status(self, link):
        try:
            response = requests.get(link)
            if response.status_code != 200:
                return link, response.url
            return None
        except requests.exceptions.RequestException:
            return link, "Connection error"
    
    def save_to_csv(self, file_path, data, header):
        with open(file_path, 'w', newline='') as csvfile:
            csv_writer = csv.writer(csvfile)
            csv_writer.writerow(header)
            csv_writer.writerows(data)

    def create_word_document_from_excel(self, broken_links_report, output_file):
        document = WordDocument()

        # Add a title to the document
        document.add_heading("Broken Links Report", level=1)

        # Iterate through the broken links and add content to the document
        for link_info in broken_links_report:
            link, page_link, section, relative = link_info
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"• On {link} there is a broken link on section '{section}', the broken link is '{page_link}'.")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

            # Add a bullet point after each link
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0.5)  # Adjust the left indentation for the bullet point
            paragraph_format.space_after = Pt(12)  # Add spacing after each paragraph

        # Save the document to the specified file path
        document.save(output_file)        

    def get_files_in_excel_folder(self):
        excel_folder = "Documents/Excel_Files"
        return [file for file in os.listdir(excel_folder) if file.endswith(".xlsx")]

        #Identify whether the link is on navigation, footer, or main content (Add left navigation)
    def identify_section(self, link, page_link, page_style):
        try:
            absolute_page_link = urljoin(link, page_link)
            response = requests.get(link)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                # Search for partial matches of the href attribute with page_link
                
                element = soup.find('a', href=lambda href: page_link in href)
                
                if element:
                    # Search for the outer elements to determine the section
                    #change to id siteheader
                    #UC-FULL-WIDTH, CONTENT-NOCAROUSEL, DYNAMICPAGE
                    #FLEXIBLECONTENT
                    #FULLWIDTH
                    #NEWANDSOCIALCOLUMNS
                    #PEOPLE
                    if page_style == "No Carousel":
                        nav_element = element.find_parent('nav')
                        if nav_element:
                            return "Navigation"
                        
                        footer_element = element.find_parent('footer')
                        if footer_element:
                            return "Footer"
                        
                        sidr_container_element = element.find_parent('div', id='sidr-container')
                        if sidr_container_element:
                            return "Main Content"
                    else:
                        return "TemplateUnknown"
                #BELOW SHOULDN'T BE USED, JUST IN CASE
                return "Unknown-notfound"  # If the element is not found in any of the relevant outer elements
            else:
                return "Unknown-pageerror"  # If there was an issue fetching the page
        except requests.exceptions.RequestException:
            return "Unknown-connectionerror"  # If there was a connection error or other issues

