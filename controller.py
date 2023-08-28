import requests
import csv
import os
from bs4 import BeautifulSoup
from urllib.parse import urljoin, unquote
from docx.shared import RGBColor, Inches, Pt
from docx import Document as WordDocument
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from concurrent.futures import ThreadPoolExecutor
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException
from PIL import Image, ImageDraw
import io
import base64
import time
import urllib.parse
import urllib
import cv2
import urllib3
import requests
from docx2pdf import convert
import os
import comtypes.client
import re
import pythoncom







class Controller:
    def __init__(self):
        self.controllerCreated = True
        
    def get_final_url(self,url):
        try:
            http = urllib3.PoolManager()
            response = requests.get(url, allow_redirects=True, timeout=10)
            return response.url
        except requests.exceptions.RequestException:
            return None
            

        
    def clean_link(self, link):
        if link and not link.startswith("mailto:") and not link.startswith("#") and not link.startswith("tel:"):
            link = unquote(link).strip()
            # CHANGED TO WHILE LOOP
            while link.endswith("%20"):
                link = link[:-1]
            return link
        return None


    def get_links_from_page_concurrently(self, url):
        try:
            with ThreadPoolExecutor() as executor:
                response = requests.get(url)
                if response.status_code == 200:
                    soup = BeautifulSoup(response.content, 'html.parser')
                    links = [self.clean_link(link.get('href')) for link in soup.find_all('a', href=True)]
                    return [link for link in links if link is not None]  # Filter out None values
                else:
                    return []
        except requests.exceptions.RequestException:
            return []
    
    def check_link_status(self, link):
        try:
            response = requests.get(link)
            if response.status_code != 200:
                return link, response.url
            return None
        except requests.exceptions.RequestException:
            return link, "Connection error"
    
    def save_to_csv(self, file_path, data, header):
        with open(file_path, 'w', newline='') as csvfile:
            csv_writer = csv.writer(csvfile)
            csv_writer.writerow(header)
            csv_writer.writerows(data)

    def create_word_document_from_excel(self, broken_links_report, output_file):
        document = WordDocument()

        # Add a title to the document
        document.add_heading("Broken Links Report", level=1)

        # Iterate through the broken links and add content to the document
        for link_info in broken_links_report:

            link, page_link, section, relative, link_name = link_info
            found=False
            tries=0
            print('taking picture!')
            try:
                while tries<10 and not found:
                    photo, found = self.get_photo(link, link_name, relative,section)
                    tries+=1
            except Exception as e:
                print(f'Error is {e}')
                photo='icon/photo_no_available.png'
            
            if 'None' in photo:
                photo='icon/photo_no_available.png'
            paragraph = document.add_paragraph()
            if 'Main Navigation' in section or 'tab' in link_name or 'accordion' in link_name:
                link_name=self.word_link_cleaner(link_name)
            # Add the link as a run with hyperlink properties
            run = paragraph.add_run("• Broken link on ")
            self.add_hyperlink(paragraph,link,link)
                
            # Add the rest of the text
            run = paragraph.add_run(f" on ")
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
            
            # Add bold section text
            run = paragraph.add_run(section)
            run.bold = True
            
            run = paragraph.add_run(f" section, the broken link is displayed as: ")
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
            
            # Add bold section text
            run = paragraph.add_run(link_name)
            run.bold = True
            
            run = paragraph.add_run(f" and the link is pointing to : ")
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

            # Add bold link_name text with hyperlink
            if page_link != relative and relative.startswith("http"):
                link_display = relative  
            else:
                link_display = page_link  
            self.add_hyperlink(paragraph, page_link, link_display)


            #PUT PHOTO HERE ON WORD
            if photo:  # Check if photo is obtained
                try:
                    document.add_picture(photo, width=Inches(5))  # Add photo to document
                except Exception as e:
                    photo='icon/photo_no_available.png'
                    print('Photo not found continuing with default picture')
                    document.add_picture(photo, width=Inches(5))

            # Add a bullet point after each link
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0.5)  # Adjust the left indentation for the bullet point
            paragraph_format.space_after = Pt(12)  # Add spacing after each paragraph

        # Provide the path of the folder you want to clean
        self.clean_folder('photos')

        # Save the document to the specified file path
        document.save(output_file) 
        print('Document Saved')
        time.sleep(3)
        self.create_pdf_from_word(output_file)


    def create_pdf_from_word(self, filename):
        try:
            pythoncom.CoInitialize()  # Initialize COM library

            # Output PDF filename
            output_pdf_filename = f"{os.path.splitext(filename)[0]}.pdf"  # Using the same name as the Word file, but with PDF extension

            # Initialize COM and Word
            word = comtypes.client.CreateObject('Word.Application')
            try:
                doc = word.Documents.Open(os.path.abspath(filename))
            except Exception as e:
                print(f"An error occurred: {e}")
            # Convert Word to PDF
            doc.SaveAs(os.path.abspath(output_pdf_filename), FileFormat=17)

            # Print the full path of the saved PDF file
            pdf_full_path = os.path.abspath(output_pdf_filename)
            print(f"Word document '{filename}' converted to PDF and saved as '{pdf_full_path}'.")
        except Exception as e:
            print(f"An error occurred: {e}")
        finally:
            # Close the document and quit Word
            doc.Close()
            word.Quit()

    def word_link_cleaner(self, link):
        if 'http' in link:
            return link
        else:
            try:
                links = link.split('/')
            except Exception as e:
                pass
            return links[-1]      

    def find_first_number_index(self, s):
        match = re.search(r'\d', s)
        if match:
            return match.end()
        return -1
    
    def get_photo(self, link, link_name, relative_link, section):
                # Initialize WebDriver and navigate to the webpage
        
        driver = webdriver.Chrome()
        driver.maximize_window()  # Open the WebDriver in full screen mode

        driver.get(f"{link}")

        if link_name == 'None' or link_name == None:
            link_name=''

        # Search for the sequence of tabs/dropdowns
        search_sequence = link_name
        print(f'Getting photo, link name is: {link_name}')
        # Check the conditions before splitting
        rest_sequence=search_sequence
        clicks=[search_sequence]
        if not search_sequence.startswith('http') and ('Main Navigation' in section or 'accordion' in link_name or 'tab' in link_name):
            print('splitting')
            try:
                print(search_sequence)
                first_number_index = self.find_first_slash_index(search_sequence)
                print(first_number_index)
                if first_number_index != -1:
                    tab_or_accordion = search_sequence[:first_number_index]
                    rest_sequence = search_sequence[first_number_index:]
                    if rest_sequence.startswith('/http'):
                        rest_sequence = rest_sequence[1:]
                    clicks = [tab_or_accordion, rest_sequence]
                    print(f"Splits are: {clicks}")
                    print(f'Sequence is: {rest_sequence}')
                else:
                    # Handle the case where there's no split or only one part
                    clicks = [search_sequence]
            except Exception as e:
                print(f"An error occurred: {e}")
        if rest_sequence.startswith('/'):
            search_word = rest_sequence[1:]
        else:
            search_word=rest_sequence

        #target_href = "accordion3"
        if "'" in search_word:
            search_word = search_word.split("'")[0]

        # Clean up the search_word
        search_word = search_word.lstrip(" =-,.!?").strip()


        # Loop through the clicks, excluding the last one
        while len(clicks) > 1:  # Exclude the last element
            print('Tab or accordion detected...')
            click = clicks.pop(0)  # Get the next click from the list
            driver.execute_script("window.scrollTo(0, 0);")
            while True:
                try:
                    print(click)
                    # Find and click on the tab or dropdown element
                    if 'accordion' in click or 'tab' in click:
                        click_element = driver.find_element(By.CSS_SELECTOR, f"a[href='#{click}']")
                    else:
                        click_element = driver.find_element(By.XPATH, f"//a[contains(text(), '{click}')]")

                    
                    # Scroll the element into view using JavaScript
                    driver.execute_script("arguments[0].scrollIntoView();", click_element)

                    # Give some time for the page to adjust after scrolling
                    time.sleep(1)

                    # Scroll a bit more to ensure proper visibility
                    driver.execute_script("window.scrollBy(0, -400);")
                    click_element.click()
                    
                    break
                except NoSuchElementException:
                    print(f"Element '{click}' not found. Moving to the next step.")
                    # Get the current scroll height
                    scroll_height = driver.execute_script("return document.body.scrollHeight")

                    # Scroll down
                    driver.find_element(By.TAG_NAME, "body").send_keys(Keys.END)

                    # Give some time for content to load (you may need to adjust this)
                    time.sleep(1)

                    # Check if you've reached the end
                    if scroll_height == prev_scroll_height:
                        print("Reached the end of the page.")
                        break

                    # Update the previous scroll height
                    prev_scroll_height = scroll_height
        
        driver.execute_script("window.scrollTo(0, 0);")
        # Define a variable to keep track of previous scroll height
        prev_scroll_height = 0
        photo_name = search_word
        photo='photo_no_available.png'
        #if search_word.startswith('http') or search_word.startswith('/') or search_word == '' or search_word == ' ' or 'tab' in search_word or 'accordion'  in search_word':
        if '/' in search_word or search_word == '' or search_word == ' ':
            photo_name='photo_with_link'
        # Scroll loop
        element_found=False
        while True:
            print('Taking picture')
            try:

                #print(f'Search word is {search_word}')
                #element = driver.find_element(By.XPATH, f"//*[contains(text(), '{search_word}')]")
                #print(f'element all is: {element}')

                #element = driver.find_element(By.XPATH, f"//a[contains(text(), '{search_word}')]")
                element=None
                elements1=[]
                elements2=[]
                try:
                    print(f'Search word is: "{search_word}"')
                    if search_word != '' or search_word != None:
                        elements1 = driver.find_elements(By.XPATH, f"//a[contains(text(), '{search_word}')]")
                    
                    saved_link=relative_link
                    relative_href=relative_link
                    if 'Main Navigation' in section:
                        relative_href = relative_link.split('/')[-1]  # Extract the fragment from the full URL
                    if relative_href is None or relative_href=='':
                        relative_href=saved_link
                    elements2 = driver.find_elements(By.CSS_SELECTOR, f"a[href*='{relative_href}']")
                    
                    if not elements2:
                        modified_string = relative_href.replace(" ", "%20")
                        elements2 = driver.find_elements(By.CSS_SELECTOR, f"a[href*='{modified_string}']")
                    
                    print(f'Elements 2: {elements2}')
                    common_elements = set(elements1) & set(elements2)
                    common_elements_list = list(common_elements)
                    print("Elements present in both lists:")
                    print(common_elements_list)
                    if common_elements_list:
                        element=common_elements_list[0]
                except Exception as e:
                    print(f'Error is {e}')
                
                print(element)
                print(search_word)
                print(relative_href)
                print(f'Elements 1: {elements1}')
                print(f'Elements 2: {elements2}')
                if element is None:
                    element_is_found=False
                    for element2 in elements2:
                        print('looking for element in elements2')
                        if search_word in element2.get_attribute("outerHTML") and not element_is_found:
                            element = element2
                            element_is_found=True
                            break
                        if element2:
                            try:
                                element = element2.find_element(By.TAG_NAME, 'img')
                                print("Element contains a child <img> tag.")
                                break
                            except NoSuchElementException:
                                print("Element does not contain a child <img> tag.")
                    if not element_is_found:
                        for element1 in elements2:
                            if relative_href in element1.get_attribute("outerHTML") and not element_is_found:
                                element = element1
                                element_is_found=True
                    if not element_is_found:
                        element = driver.find_element(By.XPATH, f"//a[contains(text(), '{search_word}')]")
                    
                
                #if len(click_elements) == 1:
                #    click_element = click_elements[0]
                #else:
                #    click_element = driver.find_element(By.XPATH, f"//a[contains(text(), '{click}') and @href='#{relative_link}']")
                #    print(f"The click element is then: {click_element}")
                if element is None:
                    parts = search_word.split('/')
                    common_elements = []

                    for part in parts:
                        common_elements_part = driver.find_elements(By.XPATH, f"//a[contains(@href, '{part}')]")
                        if common_elements_part:
                            if not common_elements:
                                common_elements = common_elements_part
                            else:
                                common_elements = [element for element in common_elements if element in common_elements_part]

                    if common_elements:
                        element = common_elements[0]  # Returning the first common element found
                    else:
                        element = None
                        print('Element not found in page')

                print(f'The element is: {element.get_attribute("outerHTML")}')
                
                # Get location and size of element
                location = element.location
                size = element.size
                locy=location['y']

                print(location)
                print(size)
                if location['y']>0 and location['x']>0:
                    element_found=True
                elif 'small navigation' in section:
                    location['y']=100
                    location['x']=200
                    size['width']=400
                    size['height']=100

                window_size = driver.get_window_size()

                # Extract width and height from the window size dictionary
                window_width = window_size['width']
                window_height = window_size['height']
                
                page_height = driver.execute_script("return document.body.scrollHeight")
                
                driver.execute_script(f"window.scrollBy(0, {location['y']});")
                
                time.sleep(3)
                closeness=page_height-locy
                print(f'Location of Element: {locy}')
                print(f'closeness: {closeness}')

                if closeness<window_height:
                    print('Way too down')
                    
                    scrollAmmount=(400-(window_height-closeness))+120
                    print(f'scrollAmmount: {scrollAmmount}')
                    driver.execute_script(f"window.scrollBy(0, -{scrollAmmount});")
                else:
                    print("Not too down")
                    driver.execute_script(f"window.scrollBy(0, -400);")

                

                # Scroll a bit more to ensure proper visibility
                #driver.execute_script(f"window.scrollBy(0, -400);")
                
                
                
                # Give some time for the page to adjust after scrolling
                time.sleep(1)
                
                 # Define screenshot area around element
                if location['y']<400:
                    y1=int(location['y'] * 1.20)
                    y2=int(y1 + size['height'] * 2.5)
                else:
                    y1=400
                    y2=600
                    
                # Define screenshot area around element
                if location['y']<200:
                    y=100
                    h=600
                else:
                    y = 150
                    h = 600
                    
                
                

                scroll_height = driver.execute_script("return document.body.scrollHeight")

                # Calculate the amount scrolled
                scrolled_amount = scroll_height - prev_scroll_height
                print('Saving rect')
               
                # Capture the screenshot using Selenium
                screenshot_base64 = driver.get_screenshot_as_base64()
                screenshot_bytes = base64.b64decode(screenshot_base64)
                pil_image_with_rect = Image.open(io.BytesIO(screenshot_bytes))
                print('Saved rect')
                # Draw a rectangle on the specified area
                draw = ImageDraw.Draw(pil_image_with_rect)
                x1 = int(location['x'] * 1.20)
                x2 = int(x1 + size['width'] * 1.60)+50
                draw.rectangle([x1, y1, x2, y2], outline=(255, 0, 0), width=3)
                print('Drawing')
                # Save the screenshot with rectangle
                pil_image_with_rect.save(f'photos/{photo_name}_with_rect.png')

                # Crop the image
                cropped_image = pil_image_with_rect.crop((0, y, pil_image_with_rect.width, y + h))  # Crop only the specified height

                # Save the cropped image
                cropped_image.save(f"photos/{photo_name}.png")

                # Print the scrolled amount
                print(f"Scrolled amount: {scrolled_amount}")

                print(element.location)
                print(element.size)

                print("Screenshot of the element captured with rectangle.")

                break
            except NoSuchElementException:
                print("Element not found. Scrolling...")
            
            except Exception as e:
                print(f"The error is {e}")
                
            # Get the current scroll height
            scroll_height = driver.execute_script("return document.body.scrollHeight")
            
            # Scroll down
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.END)
            
            # Give some time for content to load (you may need to adjust this)
            time.sleep(2)
            
            # Check if you've reached the end
            if scroll_height == prev_scroll_height:
                print("Reached the end of the page.")
                break
            
            # Update the previous scroll height
            prev_scroll_height = scroll_height
            
        # Close the WebDriver
        driver.quit()

        return f"photos/{photo_name}.png", element_found

    def find_first_slash_index(self, s):
        print(f'slash is: {s}')
        match = re.search(r'/', s)

        if match:
            return match.start()
        return -1



    def get_files_in_excel_folder(self):
        excel_folder = "Documents/Excel_Files"
        return [file for file in os.listdir(excel_folder) if file.endswith(".xlsx")]

        #Identify whether the link is on navigation, footer, or main content (Add left navigation)
    def identify_section(self, link, page_link, page_style):
        try:
            absolute_page_link = urljoin(link, page_link)
            response = requests.get(link)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                # Search for partial matches of the href attribute with page_link
                page_link=urllib.parse.unquote(page_link)  # Remove any URL encoding

                delimiter_characters = ['~', '%0D', '%']
                common_elements = []
                saved_link=page_link
                for delimiter in delimiter_characters:
                    if delimiter in page_link:

                        parts = page_link.split(delimiter)
                        try:
                            elements_part = [soup.find_all('a', href=lambda href: href is not None and part in href.strip()) for part in parts]

                        except Exception as e:
                            print('Error on First')
                        print(f'Element parts are: {elements_part}')
                        if elements_part:
                            if not common_elements:
                                common_elements = elements_part
                            else:
                                common_elements = [element for element in common_elements if element in elements_part]
                        print(common_elements)
                    if common_elements:
                        element = common_elements[0][0]  # Returning the first common element found
                        break
                    else:
                        element = soup.find('a', href=lambda href: href is not None and href.strip() == saved_link)

                print('~~~~~~~~~~~~~~~~~~~')
                print(saved_link)
                if element is None:
                    print('here Inside')
                    print(saved_link)
                    for delimiter in delimiter_characters:
                        if delimiter in saved_link:

                            parts = saved_link.split(delimiter)

                            try:
                                elements_part = [soup.find_all('a', href=lambda href: href is not None and part in href.strip()) for part in parts]

                            except Exception as e:
                                print('Error on Second')
                            print(f'Element parts are: {elements_part}')
                            if elements_part:
                                if not common_elements:
                                    common_elements = elements_part
                                else:
                                    common_elements = [element for element in common_elements if element in elements_part]
                        
                        if common_elements:
                            if common_elements[0]:
                                element = common_elements[0][0]  # Get the link from the second inner list
                            elif common_elements[1]:
                                element = common_elements[1][0]  # Get the link from the first inner list
                        else:
                            element = soup.find('a', href=lambda href: href is not None and href.strip() == saved_link)

                        if element:
                            break
                    print(f'Elemento is: {element}')

                if element:
                    element_name = 'No_name'
                    element_name = element.get_text(strip=True)
                    
                    # Check for accordion or tab
                    accordion_div = element.find_parent('div', class_='collapsable collapse')
                    if accordion_div:
                        accordion_id = accordion_div.get('id')
                        if accordion_id:
                            element_name = f"{accordion_id}/{element_name}"

                    tab_div = element.find_parent('div', class_='tab-pane')
                    if tab_div:
                        tab_id = tab_div.get('id')
                        if tab_id:
                            element_name = f"{tab_id}/{element_name}"


                    if page_style == "No Carousel":

                        left_nav_element = element.find_parent('div', class_='leftSidebar')
                        if left_nav_element:
                            return "Left Navigation", element_name

                        main_element = element.find_parent('div', class_='rightSidebar')
                        if main_element:
                            return "Main Content", element_name   
                        
                        breadcrumb = element.find_parent(class_='breadcrumb')
                        if breadcrumb:
                            return "Main Content (Top small navigation)", element_name
                        
                        nav_element = element.find_parent('nav')
                        if nav_element:
                            # Find the outer ul element with class "nav navbar-nav" within the nav_element
                            outer_ul = nav_element.find('ul', class_='nav navbar-nav')

                            if outer_ul and 'navbar-right' not in outer_ul.get('class', []):
                                # Loop through <li> elements with class 'dropdown'
                                for li in outer_ul.find_all('li', class_='dropdown'):
                                    # Check if the element is a descendant of the current <li>
                                    if element in li.descendants:
                                        # Find the specific <a> element within the <li>
                                        dropdown_a = li.find('a', class_='dropdown-toggle')
                                        if dropdown_a:
                                            dropdown_name = dropdown_a.get_text(strip=True)
                                            return f"Main Navigation, tab name: {dropdown_name}", f"{dropdown_name}/{element_name}"
                                    
                                #return "Main Navigation, tab name: {}", element_name

                        resources_li = element.find_parent('li', class_='dropdown resourcesForLinks')
                        if resources_li:
                            return "Main Navigation, tab name: Resources For Links", f"Resources For/{element_name}"

                        # If not found in dropdowns, check in "Quick Links"
                        quick_links_li = element.find_parent('li', class_='hidden-below-1365')
                        if quick_links_li:
                            return "Main Navigation, tab name: Quick Links", f"Quick Links/{element_name}"

                        footer_element = element.find_parent('footer')
                        if footer_element:
                            return "Footer", element_name

                        #Fullscreen
                        section_anchor_div = soup.find('div', class_='sectionAnchor')
                        if section_anchor_div:
                            section_id = section_anchor_div.get('id')
                            
                            if section_id:
                                return f"Main Content: Section ID: {section_id}" , element_name

                            else:
                                return f"Main Content: Section ID: No ID", element_name

                        
                        
                        sidr_container_element = element.find_parent('div', id='sidr-container')
                        if sidr_container_element:
                            return "General content", element_name
                    else:
                        return "TemplateUnknown","None"
                #Error case
                return "General content","None"  # If the element is not found in any of the relevant outer elements
            else:
                return "Unknown-pageerror","None"  # If there was an issue fetching the page
        except requests.exceptions.RequestException:
            return "Unknown-connectionerror","None"  # If there was a connection error or other issues



    def clean_folder(self, folder_path):
        # Check if the folder exists
        if os.path.exists(folder_path) and os.path.isdir(folder_path):
            # Iterate over all files and subdirectories in the folder
            for item in os.listdir(folder_path):
                item_path = os.path.join(folder_path, item)
                if os.path.isfile(item_path):
                    # Remove files
                    os.remove(item_path)
                elif os.path.isdir(item_path):
                    # Remove subdirectories and their contents recursively
                    clean_folder(item_path)
                    os.rmdir(item_path)
        else:
            print(f"The folder '{folder_path}' does not exist.")

    def add_hyperlink(self, paragraph, url, text):
        """
        A function that places a hyperlink within a paragraph object.
        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """
        
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        # Create a new w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.ns.qn('w:val'), 'single')  # 'single' for single underline
        rPr.append(u)

        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(docx.oxml.ns.qn('w:val'), '0000FF')  # Blue color as HEX value
        rPr.append(color)

        new_run.append(rPr)

        # Join all the xml elements together and add the required text to the w:r element
        t = docx.oxml.shared.OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)

        # Append the hyperlink to the paragraph
        paragraph._p.append(hyperlink)

        return hyperlink

    



